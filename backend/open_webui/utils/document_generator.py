import io
import logging
import os
import re
from datetime import datetime
from typing import Optional

from bs4 import BeautifulSoup
from markdown import markdown

log = logging.getLogger(__name__)

# ──────────────────────────────────────────────
# Pydantic models (Local copies for internal use)
# ──────────────────────────────────────────────


class MessageItem:
    def __init__(self, role: str, content: str, timestamp: Optional[int] = None):
        self.role = role
        self.content = content
        self.timestamp = timestamp


# ──────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────

_ROLE_RU = {
    'user': 'Пользователь',
    'assistant': 'Ассистент',
    'system': 'Система',
}


def _role_label(role: str) -> str:
    return _ROLE_RU.get(role, role.capitalize())


def _safe_filename(name: str) -> str:
    """Remove characters unsafe for filenames."""
    return re.sub(r'[^\w\s\-.]', '', name).strip() or 'document'


def _is_table_separator(line: str) -> bool:
    """Check if a line is a markdown table separator (e.g. |---|---|)."""
    line = line.strip()
    if not line or '|' not in line:
        return False
    # Remove common table separator characters and see if anything else is left
    # Standard separators: | --- | :---: | ---: |
    remaining = re.sub(r'[\|\-\s:]', '', line)
    return len(remaining) == 0


def _parse_markdown_tables(text: str) -> list[list[list[str]]]:
    """
    Extract all markdown tables from *text*.
    Returns a list of tables; each table is a list of rows;
    each row is a list of cell strings.
    """
    tables: list[list[list[str]]] = []
    # Use BeautifulSoup to find tables in converted markdown
    html = markdown(text, extensions=['tables'])
    soup = BeautifulSoup(html, 'html.parser')
    for table_tag in soup.find_all('table'):
        table_data = []
        # Header
        thead = table_tag.find('thead')
        if thead:
            for tr in thead.find_all('tr'):
                cells = [td.get_text(strip=True) for td in tr.find_all(['th', 'td'])]
                if cells:
                    table_data.append(cells)
        # Body
        tbody = table_tag.find('tbody')
        if tbody:
            for tr in tbody.find_all('tr'):
                cells = [td.get_text(strip=True) for td in tr.find_all(['th', 'td'])]
                if cells:
                    table_data.append(cells)
        elif not thead:
            # Table might not have thead/tbody
            for tr in table_tag.find_all('tr'):
                cells = [td.get_text(strip=True) for td in tr.find_all(['th', 'td'])]
                if cells:
                    table_data.append(cells)

        if table_data:
            tables.append(table_data)
    return tables


def _messages_to_plain(messages: list[MessageItem]) -> str:
    parts = []
    for msg in messages:
        parts.append(f'{_role_label(msg.role).upper()}:\n{msg.content}')
    return '\n\n---\n\n'.join(parts)


def safe_fpdf_text(text: str) -> str:
    """Best-effort Cyrillic → latin-1 for core fonts (transliteration fallback)."""
    try:
        return text.encode('latin-1').decode('latin-1')
    except (UnicodeEncodeError, UnicodeDecodeError):
        # Transliterate Cyrillic
        cyr_map = {
            'а': 'a',
            'б': 'b',
            'в': 'v',
            'г': 'g',
            'д': 'd',
            'е': 'e',
            'ё': 'yo',
            'ж': 'zh',
            'з': 'z',
            'и': 'i',
            'й': 'y',
            'к': 'k',
            'л': 'l',
            'м': 'm',
            'н': 'n',
            'о': 'o',
            'п': 'p',
            'р': 'r',
            'с': 's',
            'т': 't',
            'у': 'u',
            'ф': 'f',
            'х': 'kh',
            'ц': 'ts',
            'ч': 'ch',
            'ш': 'sh',
            'щ': 'shch',
            'ъ': '',
            'ы': 'y',
            'ь': '',
            'э': 'e',
            'ю': 'yu',
            'я': 'ya',
        }
        result = []
        for ch in text:
            lo = ch.lower()
            if lo in cyr_map:
                rep = cyr_map[lo]
                result.append(rep.upper() if ch.isupper() else rep)
            else:
                try:
                    ch.encode('latin-1')
                    result.append(ch)
                except (UnicodeEncodeError, UnicodeDecodeError):
                    result.append('?')
        return ''.join(result)


# ──────────────────────────────────────────────
# Core Generation Functions
# ──────────────────────────────────────────────


def generate_txt(title: str, messages: list[MessageItem]) -> bytes:
    header = f'{title}\nЭкспорт: {datetime.now().strftime("%d.%m.%Y %H:%M")}\n{"=" * 60}\n\n'

    parts = []
    for msg in messages:
        role_line = f'[{_role_label(msg.role).upper()}]'
        parts.append(f'{role_line}\n{msg.content}')

    body = '\n\n'.join(parts)
    return (header + body).encode('utf-8')


def generate_md(title: str, messages: list[MessageItem]) -> bytes:
    lines = [
        f'# {title}',
        '',
        f'> Дата экспорта: {datetime.now().strftime("%d.%m.%Y %H:%M")}',
        '',
        '---',
        '',
    ]
    for msg in messages:
        lines.append(f'### {_role_label(msg.role).upper()}')
        lines.append('')
        lines.append(msg.content)
        lines.append('')
        lines.append('---')
        lines.append('')

    return '\n'.join(lines).encode('utf-8')


def generate_pdf(title: str, messages: list[MessageItem]) -> bytes:
    try:
        from fpdf import FPDF

        class ChatPDF(FPDF):
            def __init__(self, *args, **kwargs):
                super().__init__(*args, **kwargs)
                # Try to load NotoSans font for Unicode/Cyrillic support
                font_path = 'backend/open_webui/static/fonts/NotoSans-Regular.ttf'
                font_bold_path = 'backend/open_webui/static/fonts/NotoSans-Bold.ttf'
                font_italic_path = 'backend/open_webui/static/fonts/NotoSans-Italic.ttf'

                if os.path.exists(font_path):
                    self.add_font('NotoSans', '', font_path)
                    self.font_name = 'NotoSans'
                else:
                    self.font_name = 'Helvetica'

                if os.path.exists(font_bold_path):
                    self.add_font('NotoSans', 'B', font_bold_path)

                if os.path.exists(font_italic_path):
                    self.add_font('NotoSans', 'I', font_italic_path)

            def header(self):
                self.set_font(self.font_name, 'B', 10)
                self.set_text_color(120, 120, 120)
                self.cell(0, 8, title or 'Экспорт чата', align='C', new_x='LMARGIN', new_y='NEXT')
                self.set_draw_color(220, 220, 220)
                self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
                self.ln(4)

            def footer(self):
                self.set_y(-15)
                self.set_font(self.font_name, 'I', 8)
                self.set_text_color(150, 150, 150)
                self.cell(0, 10, f'Стр. {self.page_no()}', align='C')

        pdf = ChatPDF(orientation='P', unit='mm', format='A4')
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()

        pdf.set_font(pdf.font_name, 'B', 16)
        pdf.set_text_color(30, 30, 30)
        pdf.multi_cell(0, 10, title or 'Экспорт чата', align='C')
        pdf.ln(2)
        pdf.set_font(pdf.font_name, '', 9)
        pdf.set_text_color(120, 120, 120)
        pdf.cell(
            0,
            6,
            f'Дата экспорта: {datetime.now().strftime("%d.%m.%Y %H:%M")}',
            align='C',
            new_x='LMARGIN',
            new_y='NEXT',
        )
        pdf.ln(6)

        role_colors = {
            'user': (26, 115, 232),
            'assistant': (13, 150, 74),
            'system': (136, 68, 0),
        }

        for i, msg in enumerate(messages):
            role = msg.role
            color = role_colors.get(role, (80, 80, 80))

            pdf.set_font(pdf.font_name, 'B', 9)
            pdf.set_text_color(*color)
            pdf.cell(0, 6, _role_label(role).upper(), new_x='LMARGIN', new_y='NEXT')

            # Render content with tables
            pdf.set_font(pdf.font_name, '', 10)
            pdf.set_text_color(40, 40, 40)

            # Split content by tables
            lines = msg.content.splitlines()
            current_text_block = []

            j = 0
            while j < len(lines):
                line = lines[j]

                # Header detection
                header_match = re.match(r'^(#{1,6})\s+(.*)', line)
                if header_match:
                    if current_text_block:
                        pdf.multi_cell(0, 5, '\n'.join(current_text_block), new_x='LMARGIN', new_y='NEXT')
                        current_text_block = []
                        pdf.ln(2)

                    level = len(header_match.group(1))
                    header_text = header_match.group(2)
                    # H1: 18, H2: 16, H3: 14, H4: 12, H5: 11, H6: 10
                    size = max(10, 20 - (level * 2))
                    pdf.set_font(pdf.font_name, 'B', size)
                    pdf.multi_cell(0, size / 2 + 2, header_text, new_x='LMARGIN', new_y='NEXT')
                    pdf.set_font(pdf.font_name, '', 10)  # Reset
                    pdf.ln(1)
                    j += 1
                    continue

                # Check for table start
                if '|' in line and j + 1 < len(lines) and _is_table_separator(lines[j + 1]):
                    # Flush current text block
                    if current_text_block:
                        pdf.multi_cell(0, 5, '\n'.join(current_text_block), new_x='LMARGIN', new_y='NEXT')
                        current_text_block = []
                        pdf.ln(2)

                    # Extract this table
                    table_lines = []
                    table_lines.append(lines[j])
                    table_lines.append(lines[j + 1])
                    j += 2
                    while j < len(lines) and '|' in lines[j]:
                        table_lines.append(lines[j])
                        j += 1

                    tables = _parse_markdown_tables('\n'.join(table_lines))
                    if tables:
                        for table in tables:
                            with pdf.table(
                                borders_layout='SINGLE_TOP_LINE',
                                cell_fill_color=(245, 245, 245),
                                cell_fill_mode='ROWS',
                                line_height=6,
                                text_align='LEFT',
                            ) as t:
                                for row in table:
                                    row_cells = t.row()
                                    for cell in row:
                                        row_cells.cell(cell)
                        pdf.ln(2)
                    continue
                else:
                    current_text_block.append(line)
                    j += 1

            if current_text_block:
                pdf.multi_cell(0, 5, '\n'.join(current_text_block), new_x='LMARGIN', new_y='NEXT')

            pdf.ln(3)

            if i < len(messages) - 1:
                pdf.set_draw_color(220, 220, 220)
                pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
                pdf.ln(4)

        return pdf.output()
    except Exception as e:
        log.exception(f'Error generating PDF: {e}')
        raise


def generate_docx(title: str, messages: list[MessageItem], include_chat_metadata: bool = True) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        doc = Document()
        for section in doc.sections:
            section.top_margin = Pt(50)
            section.bottom_margin = Pt(50)
            section.left_margin = Pt(60)
            section.right_margin = Pt(60)

        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        title_para = doc.add_heading(title or 'Экспорт чата', level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_para = doc.add_paragraph(f'Дата экспорта: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in date_para.runs:
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.font.size = Pt(10)
        doc.add_paragraph()

        if include_chat_metadata:
            _ROLE_COLOR = {
                'user': RGBColor(0x1A, 0x73, 0xE8),
                'assistant': RGBColor(0x0D, 0x96, 0x4A),
                'system': RGBColor(0x88, 0x44, 0x00),
            }

            for i, msg in enumerate(messages):
                role_para = doc.add_paragraph()
                role_para.paragraph_format.space_after = Pt(2)
                role_run = role_para.add_run(_role_label(msg.role).upper())
                role_run.bold = True
                role_run.font.size = Pt(10)
                role_run.font.color.rgb = _ROLE_COLOR.get(msg.role, RGBColor(0x55, 0x55, 0x55))

                # Process message content
                _append_content_to_docx(doc, msg.content)

                if i < len(messages) - 1:
                    sep = doc.add_paragraph()
                    sep.paragraph_format.space_before = Pt(4)
                    sep.paragraph_format.space_after = Pt(4)
                    pPr = sep._p.get_or_add_pPr()
                    pBdr = OxmlElement('w:pBdr')
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'single')
                    bottom.set(qn('w:sz'), '6')
                    bottom.set(qn('w:space'), '1')
                    bottom.set(qn('w:color'), 'E0E0E0')
                    pBdr.append(bottom)
                    pPr.append(pBdr)
        else:
            plain_content = '\n\n'.join(msg.content for msg in messages if msg.content)
            _append_content_to_docx(doc, plain_content)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        log.exception(f'Error generating DOCX: {e}')
        raise


def _append_content_to_docx(doc, content: str):
    import re

    from docx.shared import Pt

    # Split content into blocks of text and tables
    lines = content.splitlines()
    j = 0
    while j < len(lines):
        line = lines[j]
        # Table detection
        if '|' in line and j + 1 < len(lines) and _is_table_separator(lines[j + 1]):
            table_lines = []
            table_lines.append(lines[j])
            table_lines.append(lines[j + 1])
            j += 2
            while j < len(lines) and '|' in lines[j]:
                table_lines.append(lines[j])
                j += 1

            tables_data = _parse_markdown_tables('\n'.join(table_lines))
            for table_data in tables_data:
                if not table_data:
                    continue
                docx_table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                docx_table.style = 'Table Grid'
                for r_idx, row_data in enumerate(table_data):
                    for c_idx, cell_val in enumerate(row_data):
                        docx_table.cell(r_idx, c_idx).text = cell_val
            doc.add_paragraph()  # Spacer
        elif line.startswith('#'):
            # Header detection
            level = 0
            for char in line:
                if char == '#':
                    level += 1
                else:
                    break

            if 1 <= level <= 6:
                header_text = line[level:].strip()
                # Heading levels in docx are 0-indexed where 0 is Title, 1 is Heading 1...
                # Markdown # is usually Heading 1
                doc.add_heading(header_text, level=min(level, 9))
                j += 1
            else:
                # Not a header, just text starting with #
                doc.add_paragraph(line)
                j += 1
        else:
            # Regular line with MD support
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)

            # Very simple MD parsing for bold/italic
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)
            j += 1


def generate_xlsx(title: str, messages: list[MessageItem], include_chat_metadata: bool = True) -> bytes:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

        wb = Workbook()
        wb.remove(wb.active)

        header_font = Font(name='Calibri', bold=True, size=11, color='FFFFFF')
        header_fill_asst = PatternFill('solid', fgColor='0D964A')
        header_fill_chat = PatternFill('solid', fgColor='3C4043')
        data_font = Font(name='Calibri', size=10)
        alt_fill = PatternFill('solid', fgColor='F8F9FA')
        wrap_align = Alignment(wrap_text=True, vertical='top')
        thin = Side(style='thin', color='E0E0E0')
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        def auto_col_widths(ws, max_width: int = 80):
            for col in ws.columns:
                max_len = 0
                col_letter = col[0].column_letter
                for cell in col:
                    try:
                        # Handle multiple lines in a cell for width calculation
                        val_str = str(cell.value or '')
                        cell_max_line = max([len(line) for line in val_str.split('\n')] + [0])
                        max_len = max(max_len, cell_max_line)
                    except Exception:
                        pass
                ws.column_dimensions[col_letter].width = min(max_len + 4, max_width)

        table_idx = 0
        for msg in messages:
            # Parse all tables in the message
            tables = _parse_markdown_tables(msg.content)
            for table in tables:
                table_idx += 1
                sheet_name = f'Таблица {table_idx}'
                # Excel sheet name limit is 31 chars
                ws = wb.create_sheet(title=sheet_name[:31])
                if not table:
                    continue
                header_row = table[0]
                data_rows = table[1:]
                for col_idx, cell_val in enumerate(header_row, start=1):
                    cell = ws.cell(row=1, column=col_idx, value=cell_val)
                    cell.font, cell.fill = header_font, header_fill_asst
                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    cell.border = border
                ws.row_dimensions[1].height = 25
                for row_idx, row in enumerate(data_rows, start=2):
                    fill = alt_fill if row_idx % 2 == 0 else None
                    for col_idx, cell_val in enumerate(row, start=1):
                        cell = ws.cell(row=row_idx, column=col_idx, value=cell_val)
                        cell.font, cell.alignment, cell.border = data_font, wrap_align, border
                        if fill:
                            cell.fill = fill
                auto_col_widths(ws)
                ws.freeze_panes = 'A2'

        if include_chat_metadata:
            ws_chat = wb.create_sheet(title='Чат', index=0)
            chat_headers = ['№', 'Роль', 'Сообщение', 'Время']
            for col_idx, hdr in enumerate(chat_headers, start=1):
                cell = ws_chat.cell(row=1, column=col_idx, value=hdr)
                cell.font, cell.fill, cell.border = header_font, header_fill_chat, border
                cell.alignment = Alignment(horizontal='center', vertical='center')
            ws_chat.row_dimensions[1].height, ws_chat.freeze_panes = 22, 'A2'

            role_fills = {
                'user': PatternFill('solid', fgColor='EAF3FF'),
                'assistant': PatternFill('solid', fgColor='EAFAF1'),
                'system': PatternFill('solid', fgColor='FFF8E1'),
            }

            for row_idx, msg in enumerate(messages, start=2):
                ts_str = datetime.fromtimestamp(msg.timestamp).strftime('%d.%m.%Y %H:%M') if msg.timestamp else ''
                row_fill = role_fills.get(msg.role)
                row_data = [row_idx - 1, _role_label(msg.role), msg.content, ts_str]
                for col_idx, value in enumerate(row_data, start=1):
                    cell = ws_chat.cell(row=row_idx, column=col_idx, value=value)
                    cell.font, cell.border = data_font, border
                    cell.alignment = wrap_align if col_idx == 3 else Alignment(vertical='top')
                    if row_fill:
                        cell.fill = row_fill

            ws_chat.column_dimensions['A'].width = 6
            ws_chat.column_dimensions['B'].width = 16
            ws_chat.column_dimensions['C'].width = 80
            ws_chat.column_dimensions['D'].width = 18
            for row_idx in range(2, len(messages) + 2):
                ws_chat.row_dimensions[row_idx].height = 40

        if table_idx == 0 and not include_chat_metadata:
            ws = wb.create_sheet(title='Документ')
            ws['A1'] = title or 'Документ'
            ws['A1'].font = Font(name='Calibri', bold=True, size=12)
            ws['A3'] = '\n\n'.join(msg.content for msg in messages if msg.content)
            ws['A3'].alignment = wrap_align
            ws.column_dimensions['A'].width = 120

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()
    except Exception as e:
        log.exception(f'Error generating XLSX: {e}')
        raise


def generate_document_by_format(
    fmt: str,
    title: str,
    messages: list[MessageItem],
    include_chat_metadata: bool = True,
) -> bytes:
    """Entry point for all generation logic."""
    fmt = fmt.lower().strip('.')
    if fmt == 'txt':
        return generate_txt(title, messages)
    elif fmt == 'md':
        return generate_md(title, messages)
    elif fmt == 'pdf':
        return generate_pdf(title, messages)
    elif fmt == 'docx':
        return generate_docx(title, messages, include_chat_metadata=include_chat_metadata)
    elif fmt == 'xlsx':
        return generate_xlsx(title, messages, include_chat_metadata=include_chat_metadata)
    else:
        raise ValueError(f'Unsupported format: {fmt}')
